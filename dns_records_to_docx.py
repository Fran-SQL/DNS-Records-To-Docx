import dns.resolver
import dns.exception
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from datetime import datetime
from docx.oxml.ns import qn

def resolve_hostname_to_ip(hostname):
    try:
        answers = dns.resolver.resolve(hostname, 'A')
        for rdata in answers:
            return str(rdata)
    except (dns.resolver.NoAnswer, dns.resolver.NXDOMAIN):
        pass
    try:
        answers = dns.resolver.resolve(hostname, 'AAAA')
        for rdata in answers:
            return str(rdata)
    except (dns.resolver.NoAnswer, dns.resolver.NXDOMAIN):
        pass
    except dns.exception.Timeout:
        return "Time limit exceeded"
    except Exception:
        return "N/A"
    return "N/A"

def get_dns_records(domain):
    records = {"A": [], "AAAA": [], "MX": [], "NS": [], "TXT": [], "CNAME": []}
    record_types = list(records.keys())

    for record_type in record_types:
        try:
            answers = dns.resolver.resolve(domain, record_type)
            for rdata in answers:
                if record_type == "MX":
                    exchange = str(rdata.exchange).rstrip('.')
                    ip = resolve_hostname_to_ip(exchange)
                    records["MX"].append({"priority": rdata.preference, "exchange": exchange, "ip": ip})
                elif record_type == "NS":
                    ns_name = str(rdata).rstrip('.')
                    ip = resolve_hostname_to_ip(ns_name)
                    records["NS"].append({"name": ns_name, "ip": ip})
                else:
                    records[record_type].append(str(rdata).rstrip('.'))
        except dns.resolver.NoAnswer:
            continue
        except dns.resolver.NXDOMAIN:
            print(f"Error: Domain '{domain}' not found (NXDOMAIN).")
            return None
        except dns.exception.Timeout:
            print(f"Timeout while querying '{domain}'.")
            return None
        except Exception as e:
            print(f"Unexpected error: {e}")
            continue
    return records

def add_stylish_table(document, headers, rows):
    """Add a styled table with custom formatting."""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Light List Accent 1'
    hdr_cells = table.rows[0].cells

    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- Color de fondo corregido ---
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "2F5496")  # color azul oscuro
        hdr_cells[i]._element.get_or_add_tcPr().append(shading)

    # --- Filas ---
    for row in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ajuste opcional de ancho de columnas
    for cell in table.columns[0].cells:
        cell.width = Inches(1.5)

    return table

def create_dns_report_doc(domain, dns_data, output_filename="informe_dns.docx"):
    document = Document()

    # --- Margins ---
    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # --- Cover Page ---
    title = document.add_heading(f"DNS Report\n{domain}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(32)
    title.runs[0].font.color.rgb = RGBColor(47, 84, 150)
    document.add_paragraph("\n")

    info = document.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(datetime.now().strftime("%d %B %Y")).font.size = Pt(11)
    document.add_page_break()

    # --- Header & Footer ---
    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.text = f"DNS Report - {domain}"
    header.style = document.styles['Normal']
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer = section.footer.paragraphs[0]
    footer.text = "Confidential Report — Generated Automatically"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    # --- DNS Data ---
    document.add_heading("1. Overview", level=1)
    document.add_paragraph(f"Below are the DNS records collected for the domain {domain}. The data were collected on {datetime.now().strftime('%d %B %Y, %H:%M')}.")
    document.add_paragraph()

    order = ["A", "AAAA", "MX", "NS", "TXT", "CNAME"]
    for idx, record_type in enumerate(order, start=2):
        data = dns_data.get(record_type)
        if not data:
            continue

        document.add_heading(f"{idx}. {record_type} Records", level=1)
        document.add_paragraph()

        if record_type == "MX":
            rows = [[r["priority"], r["exchange"], r["ip"]] for r in data]
            add_stylish_table(document, ["Priority", "Exchange", "IP"], rows)
        elif record_type == "NS":
            rows = [[r["name"], r["ip"]] for r in data]
            add_stylish_table(document, ["Name Server", "IP Address"], rows)
        elif record_type in ["A", "AAAA", "CNAME", "TXT"]:
            for item in data:
                p = document.add_paragraph(style='List Bullet')
                p.add_run(item)
        document.add_paragraph()

    # --- Save ---
    document.save(output_filename)
    print(f"Report successfully generated: {output_filename}")

# --- Run ---
if __name__ == "__main__":
    target_domain = input("Enter domain to analyze (e.g. example.com): ").strip()
    if not target_domain:
        print("No domain entered. Exiting.")
    else:
        print(f"Collecting DNS data for {target_domain}...")
        data = get_dns_records(target_domain)
        if data:
            filename = f"DNS_Report_{target_domain.replace('.', '_')}.docx"
            create_dns_report_doc(target_domain, data, filename)
        else:
            print("Failed to collect DNS data.")
