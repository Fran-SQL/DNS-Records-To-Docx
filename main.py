import dns.resolver
import dns.exception
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def resolve_hostname_to_ip(hostname):
    """
    Resolves a hostname to its IPv4 or IPv6 address.
    Returns the first IP found or ‘N/A’ if it cannot be resolved.
    """
    try:
        # Try resolving A (IPv4)
        answers = dns.resolver.resolve(hostname, 'A')
        for rdata in answers:
            return str(rdata)
    except (dns.resolver.NoAnswer, dns.resolver.NXDOMAIN):
        pass # It doesn't have an A record, we tried AAAA.

    try:
        # Trying to resolve AAAA (IPv6)
        answers = dns.resolver.resolve(hostname, 'AAAA')
        for rdata in answers:
            return str(rdata)
    except (dns.resolver.NoAnswer, dns.resolver.NXDOMAIN):
        pass # It doesn't have an AAAA record.

    except dns.exception.Timeout:
        return "Time limit exceeded"
    except Exception:
        return "N/A" # Other errors

    return "N/A" # If no IP was found

def get_dns_records(domain):
    """
    Obtains various types of DNS records for a given domain, including NS and MX IPs.
    """
    records = {
        "A": [],    # Host / IPv4
        "AAAA": [], # Host / IPv6
        "MX": [],   # Mail Exchanger
        "NS": [],   # Name Servers
        "TXT": [],  # Text records
        "CNAME": [],# Canonical Name
    }

    main_record_types = ["A", "AAAA", "MX", "NS", "TXT", "CNAME"]

    for record_type in main_record_types:
        try:
            answers = dns.resolver.resolve(domain, record_type)
            for rdata in answers:
                if record_type == "MX":
                    exchange_hostname = str(rdata.exchange).rstrip('.')
                    exchange_ip = resolve_hostname_to_ip(exchange_hostname)
                    records[record_type].append({
                        "priority": rdata.preference,
                        "exchange": exchange_hostname,
                        "ip": exchange_ip
                    })
                elif record_type == "NS":
                    ns_hostname = str(rdata).rstrip('.')
                    ns_ip = resolve_hostname_to_ip(ns_hostname)
                    records[record_type].append({
                        "name": ns_hostname,
                        "ip": ns_ip
                    })
                else:
                    records[record_type].append(str(rdata).rstrip('.'))
        except dns.resolver.NoAnswer:
            pass
        except dns.resolver.NXDOMAIN:
            print(f"Error: Domain '{domain}' not found (NXDOMAIN).")
            return None
        except dns.exception.Timeout:
            print(f"Error: Timeout while querying '{domain}'.")
            return None
        except Exception as e:
            print(f"Unexpected error while getting records {record_type} from {domain}: {e}")
            pass

    # Remove duplicates and sort some types of records
    for key in ["A", "AAAA", "TXT", "CNAME"]:
        records[key] = sorted(list(set(records[key])))
    # NS and MX are dictionaries, the set doesn't work directly, we handle it separately.
    # Removal of duplicates for NS and MX (based on the name/exchange for NS, on the priority-exchange pair for MX)
    unique_ns = []
    seen_ns_names = set()
    for ns_rec in records["NS"]:
        if ns_rec["name"] not in seen_ns_names:
            unique_ns.append(ns_rec)
            seen_ns_names.add(ns_rec["name"])
    records["NS"] = sorted(unique_ns, key=lambda x: x["name"])

    unique_mx = []
    seen_mx_tuples = set()
    for mx_rec in records["MX"]:
        mx_tuple = (mx_rec["priority"], mx_rec["exchange"])
        if mx_tuple not in seen_mx_tuples:
            unique_mx.append(mx_rec)
            seen_mx_tuples.add(mx_tuple)
    records["MX"] = sorted(unique_mx, key=lambda x: x["priority"])

    return records

def create_dns_report_doc(domain, dns_data, output_filename="informe_dns.docx"):
    """
    Create a docx document with the DNS data.
    """
    document = Document()

    # --- Document Configuration ---
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # --- Main Title ---
    title_run = document.add_heading(level=0).add_run(f"DNS report for {domain}")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(28)
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()

    # --- Generation Date ---
    date_paragraph = document.add_paragraph()
    date_paragraph.add_run(f"Generated on {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}").font.size = Pt(10)
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph()

    # --- DNS Record Sections ---
    record_order = ["A", "AAAA", "MX", "NS", "TXT", "CNAME"]
    for record_type in record_order:
        data_list = dns_data.get(record_type)
        if data_list: # Only add section if there is data
            document.add_heading(f"{record_type} records", level=1) # Section Title

            if record_type == "MX":
                # Table for MX records (priority, host, IP)
                table = document.add_table(rows=1, cols=3) # 3 columns
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Priority'
                hdr_cells[1].text = 'MX server'
                hdr_cells[2].text = 'IP'
                for cell in hdr_cells:
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                for item in data_list:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(item["priority"])
                    row_cells[1].text = item["exchange"]
                    row_cells[2].text = item["ip"]
                    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                document.add_paragraph()

            elif record_type == "NS":
                # List for NS records with their IPs
                for item in data_list:
                    p = document.add_paragraph(style='List Bullet')
                    p.add_run(f"{item['name']} ({item['ip']})")
                document.add_paragraph()

            elif record_type in ["A", "AAAA", "CNAME", "TXT"]:
                # List with bullet points for other types of records
                for item in data_list:
                    p = document.add_paragraph(style='List Bullet')
                    p.add_run(item)
                document.add_paragraph()
            else:
                for item in data_list:
                    document.add_paragraph(str(item))
                document.add_paragraph()

    # --- Save the docx ---
    try:
        document.save(output_filename)
        print(f"Document '{output_filename}' successfully generated.")
    except Exception as e:
        print(f"Error saving the document: {e}")

# --- Use of the Script ---
if __name__ == "__main__":
    target_domain = input("Enter the domain to analyze (e.g., google.com): ")

    if not target_domain:
        print("No domain was entered. Exiting.")
    else:
        print(f"Obtaining DNS records for {target_domain}...")
        dns_data = get_dns_records(target_domain)

        if dns_data:
            output_filename = f"DNS_records_{target_domain.replace('.', '_')}.docx"
            create_dns_report_doc(target_domain, dns_data, output_filename)
        else:
            print("The DNS data could not be obtained or the domain is not valid/reachable.")
